import logging

from docx import Document
from docx.shared import Pt
from datetime import date
from os import path


def write_to_docx(weekly_update: str, week_start_date: str, current_directory: str):
    log = logging.getLogger()
    log.debug("Writing to .docx file")
    template_path = path.join(current_directory,"document_handler/template.docx")

    summary_file = Document(template_path)

    summary_file.add_paragraph(
        "Weekly Update: {}".format(week_start_date), style="Weekly Update Header"
    )

    for line in weekly_update.splitlines():
        line = line.removesuffix("\n")
        if line.startswith(">") or line.startswith("#"):
            line_data = line.split(sep=None, maxsplit=1)
        else:
            line_data = ["", line.strip()]

        match line_data[0]:
            case "#":
                style = "Weekly Update Header"
            case "##":
                style = "Weekly Update Sub Heading"
            case ">":
                style = "List Bullet"
            case ">>":
                style = "List Bullet 2"
            case ">>>":
                style = "List Bullet 3"
            case ">>>>":
                style = "List Bullet 4"
            case _:
                style = "Normal"

        p = summary_file.add_paragraph(line_data[1], style=style)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

    title = "Weekly_Update_{}.docx".format(week_start_date)

    summary_file.save(title)

    print("\n\nSummary file saved to {}\n\n".format(path.join(path.abspath(path.curdir), title)))
